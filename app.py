import streamlit as st
from datetime import datetime, timedelta
import requests
from docx import Document
import io

# ========================
# CẤU HÌNH
# ========================
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
GROQ_CHAT_URL = "https://api.groq.com/openai/v1/chat/completions"

st.set_page_config(page_title="Thư Ký AI PRO", layout="centered")
st.title("📑 Thư Ký AI PRO – Chuẩn Doanh Nghiệp")

# ========================
# INPUT TEXT
# ========================
meeting_text = st.text_area("Nhập nội dung cuộc họp:", height=300)

# ========================
# HÀM GỌI AI
# ========================
def generate_minutes(text):

    today = datetime.now().strftime("%d/%m/%Y")

    prompt = f"""
Bạn là Thư ký Doanh nghiệp chuyên nghiệp.

Hãy tạo biên bản họp chuẩn doanh nghiệp.

Ngày họp: {today}

Yêu cầu:
- Loại bỏ tranh luận.
- Chỉ giữ quyết định cuối.
- Chuẩn hóa KPI.
- Tạo bảng KPI.
- Tạo bảng phân công và deadline.
- Trình bày chuyên nghiệp.

Nội dung:
{text}
"""

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "llama3-70b-8192",
        "messages": [
            {"role": "system", "content": "Bạn là thư ký doanh nghiệp chuyên nghiệp."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2
    }

    response = requests.post(GROQ_CHAT_URL, headers=headers, json=data)

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        return f"Lỗi API: {response.text}"

# ========================
# TẠO BIÊN BẢN
# ========================
if st.button("🚀 Tạo Biên Bản Chuẩn Doanh Nghiệp") and meeting_text:

    with st.spinner("Đang xử lý..."):
        result = generate_minutes(meeting_text)

    st.subheader("📄 Biên bản hoàn chỉnh")
    st.write(result)

    # ========================
    # XUẤT DOCX
    # ========================
    doc = Document()
    doc.add_heading("BIÊN BẢN HỌP", level=1)
    doc.add_paragraph(result)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        "⬇️ Tải file DOCX",
        buffer,
        file_name="Bien_Ban_Hop_Doanh_Nghiep.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )